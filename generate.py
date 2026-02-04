import os
from openai import OpenAI
from dotenv import load_dotenv
import requests
from docx import Document
from docx.shared import Pt
import tempfile

load_dotenv()

client = OpenAI(
    api_key=os.getenv("AKASH_API_KEY"),
    base_url="https://chatapi.akash.network/api/v1"
)

AKASH_API_KEY = os.getenv("AKASH_API_KEY")
AKASH_URL = "https://chatapi.akash.network/v1/chat/completions"
MODEL = "Meta-Llama-3-1-70B-Instruct-FP8"


AKASH_API_KEY = os.getenv("AKASH_API_KEY")
AKASH_URL = "https://api.akashml.com/v1/chat/completions"  
SHORT_FORM_MODEL = "deepseek-ai/DeepSeek-V3.2"             
BLOG_MODEL = "deepseek-ai/DeepSeek-V3.2"                 

HEADERS = {
    "Authorization": f"Bearer {AKASH_API_KEY}",
    "Content-Type": "application/json"
}


def generate_short_form(transcript: str):
    payload = {
        "model": SHORT_FORM_MODEL,
        "messages": [
            {
                "role": "system",
                "content": "You are a social media content expert. Turn transcripts into engaging short-form posts (60–100 words)."
            },
            {
                "role": "user",
                "content": f"Convert this transcript into 3 engaging short-form posts. Number them 1, 2, and 3.\n\nTranscript:\n{transcript}"
            }
        ],
        "max_tokens": 1000
    }

    response = requests.post(AKASH_URL, headers=HEADERS, json=payload)
    
    if response.status_code != 200:
        print("Akash API Error:", response.status_code, response.text)
        raise Exception("Akash request failed")
    
    data = response.json()

    posts = []
    for i, choice in enumerate(data["choices"], 1):
        text = choice["message"]["content"].strip()
        posts.append(f"📱 Post {i}:\n{text}")

    return "\n\n" + "\n\n" + ("-"*80 + "\n\n").join(posts)

def generate_x_posts(transcript: str):
    payload = {
        "model": SHORT_FORM_MODEL,
        "messages": [
            {
                "role": "system",
                "content": (
                    "You are an expert X (Twitter) content creator. "
                    "Write 3 separate viral-style posts. "
                    "Each post must be under 280 characters, use short punchy lines, "
                    "and end with a question or strong closing statement.\n\n"
                    "IMPORTANT: Separate each post using ONLY this line:\n###\n"
                )
            },
            {
                "role": "user",
                "content": f"Create 3 X posts from this transcript:\n\n{transcript}"
            }
        ],
        "max_tokens": 800
    }

    response = requests.post(AKASH_URL, headers=HEADERS, json=payload)

    if response.status_code != 200:
        raise Exception(f"Akash request failed: {response.text}")

    data = response.json()
    raw_text = data["choices"][0]["message"]["content"].strip()

    # Split posts cleanly
    split_posts = [p.strip() for p in raw_text.split("###") if p.strip()]

    formatted_posts = []
    for i, post in enumerate(split_posts, 1):
        formatted_posts.append(f"🐦 X Post {i}:\n{post}")

    return "\n\n" + ("\n\n" + "-"*60 + "\n\n").join(formatted_posts)


def generate_linkedin_posts(transcript: str):
    payload = {
        "model": SHORT_FORM_MODEL,
        "messages": [
            {
                "role": "system",
                "content": (
                    "You are a professional LinkedIn content strategist. "
                    "Write thoughtful, story-driven LinkedIn posts that provide value, lessons, or insights. "
                    "Use short paragraphs for readability. Keep a professional but human tone. "
                    "End with a reflective takeaway or discussion-inviting question. "
                    "Each post should be 120–250 words."
                )
            },
            {
                "role": "user",
                "content": (
                    f"Turn this transcript into 2 high-value LinkedIn posts. "
                    f"Clearly separate them with 'POST 1:' and 'POST 2:'.\n\n"
                    f"Transcript:\n{transcript}"
                )
            }
        ],
        "max_tokens": 1200
    }

    response = requests.post(AKASH_URL, headers=HEADERS, json=payload)

    if response.status_code != 200:
        print("Akash API Error:", response.status_code, response.text)
        raise Exception("Akash request failed")

    data = response.json()
    full_text = data["choices"][0]["message"]["content"].strip()

    # Split posts using markers from the AI
    parts = full_text.split("POST 2:")

    post1 = parts[0].replace("POST 1:", "").strip()
    post2 = parts[1].strip() if len(parts) > 1 else ""

    formatted_output = (
        "💼 LinkedIn Post 1\n\n"
        f"{post1}\n\n"
        "--------------------------------------------------\n\n"
        "💼 LinkedIn Post 2\n\n"
        f"{post2}"
    )

    return formatted_output

def generate_blog_post(transcript: str):
    payload = {
        "model": BLOG_MODEL,  # e.g. "deepseek-ai/DeepSeek-V3.2"
        "messages": [
            {
                "role": "system",
                "content": (
                    "You are a professional blog writer. "
                    "Turn video transcripts into well-structured, engaging blog posts. "
                    "Always include:\n"
                    "• A compelling introduction\n"
                    "• Clear sections with subheadings\n"
                    "• Practical insights or examples where possible\n"
                    "• A strong conclusion\n"
                    "Write in a professional but approachable tone."
                )
            },
            {
                "role": "user",
                "content": (
                    "Convert this transcript into a blog post (800–1200 words).\n\n"
                    "Return your response in this EXACT format:\n\n"
                    "TITLE: <Engaging Blog Title>\n"
                    "META DESCRIPTION: <SEO meta description under 160 characters>\n\n"
                    "BLOG:\n"
                    "## Introduction\n"
                    "<intro text>\n\n"
                    "## Section Heading\n"
                    "<section text>\n\n"
                    "## Conclusion\n"
                    "<closing text>\n\n"
                    f"Transcript:\n{transcript}"
                )
            }
        ],
        "max_tokens": 2000,
        "temperature": 0.7
    }

    response = requests.post(AKASH_URL, headers=HEADERS, json=payload)

    if response.status_code != 200:
        print("Akash API Error:", response.status_code, response.text)
        raise Exception("Blog generation failed")

    data = response.json()
    blog_text = data["choices"][0]["message"]["content"].strip()

    return blog_text

# parsing the blog response
def parse_blog_content(raw_text: str):
    title = raw_text.split("TITLE:")[1].split("META DESCRIPTION:")[0].strip()
    meta = raw_text.split("META DESCRIPTION:")[1].split("BLOG:")[0].strip()
    blog = raw_text.split("BLOG:")[1].strip()
    return title, meta, blog

def save_blog_to_docx(blog_text: str) -> str:
    doc = Document()

    # Title
    doc.add_heading("Generated Blog Post", level=1)

    # Add paragraphs
    for line in blog_text.split("\n"):
        if line.strip() == "":
            doc.add_paragraph("")  # spacing
        else:
            para = doc.add_paragraph(line.strip())
            para.style.font.size = Pt(12)

    # Save to temp file
    temp_dir = tempfile.gettempdir()
    file_path = os.path.join(temp_dir, "blog_post.docx")
    doc.save(file_path)

    return file_path